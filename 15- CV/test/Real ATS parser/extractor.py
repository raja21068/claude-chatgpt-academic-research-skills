"""Text extraction from resume files.

Tries the high-fidelity libraries first; falls back to plain-text reading if
they are not installed.

For PDFs: uses pdfminer.six to extract layout-aware text. This is what many
real ATS platforms use as their PDF backend (or something derived from it).

For DOCX: uses python-docx to walk paragraphs and tables. This matches how
Workday/iCIMS parse DOCX files.
"""
from __future__ import annotations

from pathlib import Path
from typing import NamedTuple


class ExtractionResult(NamedTuple):
    text: str
    method: str          # 'pdfminer', 'python-docx', 'plaintext', 'markdown'
    fidelity: str        # 'high', 'medium', 'low'
    warnings: list[str]


def extract(path: Path) -> ExtractionResult:
    """Dispatch to the right extractor based on file extension."""
    ext = path.suffix.lower()
    if ext == ".pdf":
        return _extract_pdf(path)
    if ext == ".docx":
        return _extract_docx(path)
    if ext in {".md", ".markdown"}:
        return _extract_markdown(path)
    if ext in {".txt", ".text"}:
        return _extract_text(path)
    return ExtractionResult(
        text="",
        method="unknown",
        fidelity="low",
        warnings=[f"Unsupported file extension: {ext}"],
    )


def _extract_pdf(path: Path) -> ExtractionResult:
    warnings: list[str] = []
    try:
        from pdfminer.high_level import extract_text  # type: ignore
    except ImportError:
        warnings.append(
            "pdfminer.six not installed — falling back to limited extraction. "
            "Install with: pip install pdfminer.six --break-system-packages"
        )
        return ExtractionResult(
            text=f"[PDF extraction unavailable: {path.name}]",
            method="plaintext",
            fidelity="low",
            warnings=warnings,
        )

    try:
        text = extract_text(str(path)) or ""
    except Exception as e:
        warnings.append(f"pdfminer extraction failed: {e}")
        return ExtractionResult(text="", method="pdfminer", fidelity="low", warnings=warnings)

    # Common parsing-breakage heuristics
    if "\x0c" in text and text.count("\x0c") > 5:
        warnings.append(f"Detected {text.count(chr(12))} form-feed page breaks. PDF likely multi-page.")
    if text.strip() == "":
        warnings.append("PDF extracted to empty text. This usually means the PDF is image-only (scanned). ATSs WILL fail to parse this — re-export as text-PDF.")
        return ExtractionResult(text=text, method="pdfminer", fidelity="low", warnings=warnings)

    return ExtractionResult(text=text, method="pdfminer", fidelity="high", warnings=warnings)


def _extract_docx(path: Path) -> ExtractionResult:
    warnings: list[str] = []
    try:
        from docx import Document  # type: ignore
    except ImportError:
        warnings.append(
            "python-docx not installed — falling back. "
            "Install with: pip install python-docx --break-system-packages"
        )
        return ExtractionResult(
            text=f"[DOCX extraction unavailable: {path.name}]",
            method="plaintext",
            fidelity="low",
            warnings=warnings,
        )

    try:
        doc = Document(str(path))
    except Exception as e:
        warnings.append(f"python-docx failed to open: {e}")
        return ExtractionResult(text="", method="python-docx", fidelity="low", warnings=warnings)

    parts: list[str] = []
    # Walk paragraphs in order
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)

    # Walk tables — but flag them as risky
    table_count = 0
    for tbl in doc.tables:
        table_count += 1
        for row in tbl.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    if table_count > 0:
        warnings.append(
            f"DOCX contains {table_count} table(s). Workday/Taleo often scramble multi-column tables. "
            "Verify that experience/projects are NOT inside tables."
        )

    # Walk headers and footers — these are commonly stripped by ATSs
    header_text_found = False
    for section in doc.sections:
        if section.header and section.header.paragraphs:
            for p in section.header.paragraphs:
                if p.text.strip():
                    header_text_found = True
        if section.footer and section.footer.paragraphs:
            for p in section.footer.paragraphs:
                if p.text.strip():
                    header_text_found = True

    if header_text_found:
        warnings.append(
            "DOCX has content in headers/footers. Workday and Taleo frequently strip these regions. "
            "Move contact info into the document body."
        )

    text = "\n".join(parts)
    if not text.strip():
        warnings.append("DOCX extracted to empty text — file may be corrupt or use unusual structure.")
        return ExtractionResult(text=text, method="python-docx", fidelity="low", warnings=warnings)

    return ExtractionResult(text=text, method="python-docx", fidelity="high", warnings=warnings)


def _extract_markdown(path: Path) -> ExtractionResult:
    text = path.read_text(encoding="utf-8", errors="ignore")
    return ExtractionResult(
        text=text,
        method="markdown",
        fidelity="medium",
        warnings=[
            "Reading raw Markdown. ATSs see whatever export you submit — "
            "run this check against the actual PDF/DOCX you will upload."
        ],
    )


def _extract_text(path: Path) -> ExtractionResult:
    return ExtractionResult(
        text=path.read_text(encoding="utf-8", errors="ignore"),
        method="plaintext",
        fidelity="medium",
        warnings=[],
    )
